from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(r"C:\Users\11489\.openclaw\workspace-market\share\同济大学智能体交流提纲_20260617.docx")


def set_font(run, size=10.5, bold=False, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, text="", size=10.5, bold=False, color=None, align=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.runs[0]
    set_font(r, size=15 if level == 1 else 12.5, bold=True, color=(31, 78, 121))
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r)


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r)


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)

for style_name in ["Normal", "List Bullet", "List Number"]:
    style = doc.styles[style_name]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("同济大学博士线上交流提纲\nAI-agent 市场战略决策智能体开发")
set_font(r, size=18, bold=True, color=(31, 78, 121))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("时间：2026年6月17日 11:00｜用途：会前准备与打印讨论稿")
set_font(r, size=10, color=(90, 90, 90))

heading(doc, "一、交流目标")
para(
    doc,
    "本次交流不建议泛泛讨论“智能体怎么做”，而应围绕当前市场战略决策智能体的真实工程问题展开："
    "如何把已有的结构化数据、RAG、分析框架和 Agent 工作空间，组织成一个可验证、可解释、可持续演进的战略决策系统。",
)
bullet(doc, "请对方从智能体架构、任务编排、RAG 知识库、评估体系、人机协同五个角度给出建议。")
bullet(doc, "重点讨论：从“工具拼接”升级到“自主决策闭环”的路径。")
bullet(doc, "希望获得：架构把关、评估方法、工程落地建议、后续合作切入点。")

heading(doc, "二、当前项目现状简述")
bullet(doc, "workspace-market：Agent 工作空间，包含主 Agent 规则、strategy-orchestrator 规则、Skills 包装、SSE/API 兼容层、工作记忆和架构文档。")
bullet(doc, "rag-engine：真实数据与分析能力底座，包含 HybridMarketAgent、PostgreSQL 市场数据查询、PGVector/RAG、文档入库、PEST/波特五力/SWOT/4P、报告生成工具。")
bullet(doc, "已具备：结构化汽车市场数据、基础 RAG 链路、分析框架工具、Agent 架构规范。")
bullet(doc, "主要卡点：strategy-orchestrator 还没有成为运行主线；RAG 知识库内容较薄；旧 Python wrapper 仍承担部分流程逻辑；报告可信度、证据链、置信度评估还不够硬。")

heading(doc, "三、建议开场表达")
para(
    doc,
    "我们现在已经有汽车市场结构化数据库、PGVector/RAG、PEST/SWOT/波特五力等工具，也做了一个市场战略 Agent 的工作空间。"
    "现在的问题不是单点能力，而是如何把这些能力组织成一个可信的战略决策智能体：它能自己判断查什么数据、什么时候补证据、"
    "如何处理冲突、如何给出可解释的置信度。想请你们从智能体架构、评估体系和知识库建设三个角度帮我们把把关。",
    color=(70, 70, 70),
)

heading(doc, "四、建议重点提问")
questions = [
    (
        "1. 智能体的“自主编排大脑”应该怎么设计？",
        "我们现在有 SQL、RAG、PEST、SWOT、报告生成等工具，但还在从固定流程迁移到 Agent 自主编排。复杂市场分析任务中，调度 Agent 应该如何判断“先查什么、再分析什么、什么时候停止”？",
        "我的理解：不能靠固定 pipeline，也不能只靠关键词路由。较合理的是 ReAct 循环：Plan -> Act -> Observe -> Reflect -> Re-plan。每轮维护证据账本，判断证据是否足够、是否冲突、是否需要补查。",
    ),
    (
        "2. 市场战略类 Agent 的评估标准应该怎么定？",
        "对于汽车市场战略分析 Agent，不能只看回答是否流畅。应该如何评估它的质量？是事实准确率、证据覆盖率、建议可执行性，还是预测命中率？",
        "我的理解：至少要评估数据正确性、证据完整性、推理质量和决策价值。尤其需要建立 evaluation rubric，后续才能证明 Agent 有实际价值。",
    ),
    (
        "3. RAG 知识库应该如何建设，才能支撑战略决策？",
        "现在 RAG 技术链路已经通了，但知识库内容很薄。对于产业战略 Agent，RAG 文档应该如何分层、打标签、更新，才能支撑可靠分析？",
        "我的理解：RAG 不能只是把 PDF 扔进去。建议按证据类型建设：行业报告、政策文件、竞品材料、用户舆情、历史报告。Metadata 至少包括 source、publish_date、brand、car_model、category、segment。",
    ),
    (
        "4. 结构化数据和 RAG 文档如何融合？",
        "结构化销量数据和非结构化行业报告经常口径不同。Agent 应该如何融合这两类证据？冲突时应该怎么处理？",
        "我的理解：结构化数据回答“多少、排名、变化”；RAG 回答“为什么、政策背景、趋势解释”。冲突时不能强行合并，要输出冲突原因、采用口径和置信度下降说明。",
    ),
    (
        "5. 置信度应该怎么计算，不能只是模型自己说一个数？",
        "当前 Agent 会输出置信度，但更多是规则估计。实际项目里，智能体的置信度应该怎么设计才可信？",
        "我的理解：置信度应综合数据新鲜度、覆盖范围、来源可信度、多源一致性、缺失字段、人工确认、历史验证等因素。",
    ),
    (
        "6. 多 Agent 分工到底有没有必要？",
        "我们现在设想有 strategy-orchestrator、data-agent、analysis-agent、report-agent。实际工程里，多 Agent 是否必要？什么时候该多 Agent，什么时候一个 Agent 加工具就够？",
        "我的理解：简单问答一个 Agent 调工具即可。市场战略任务更适合多 Agent，但关键不是“多”，而是职责清晰、可测试、可恢复。",
    ),
    (
        "7. 如何让 Agent 具备长期记忆和复盘能力？",
        "市场战略 Agent 需要记住过去判断、错误、用户偏好和历史报告。长期记忆应该怎么设计？哪些内容进 memory，哪些进 RAG，哪些进数据库？",
        "我的理解：memory 记录任务、偏好、架构决策和错误；RAG 存报告、政策、行业文档；数据库存销量、配置等结构化指标；evidence ledger 存每次分析的证据链。",
    ),
    (
        "8. 人机协同边界怎么设？",
        "战略决策不能完全交给 AI。哪些环节应该让 Agent 自主完成，哪些环节必须人工确认？",
        "我的理解：Agent 可自主完成数据收集、初步分析、竞品对比、风险扫描和报告草稿；关键战略建议、高投入决策、数据口径切换、低置信度结论应保留人工确认。",
    ),
]

for title_text, ask, answer in questions:
    heading(doc, title_text, 2)
    para(doc, "建议提问：", bold=True)
    para(doc, ask)
    para(doc, "我的理解与预期回答方向：", bold=True)
    para(doc, answer)

heading(doc, "五、最值得现场重点追问的 3 个问题")
number(doc, "如何设计战略决策 Agent 的评估体系？这是后续证明 Agent 价值的关键。")
number(doc, "如何让结构化数据和 RAG 证据融合，并处理冲突？这是当前项目最核心的技术难点。")
number(doc, "复杂任务的自主编排应该怎么实现，是否需要多 Agent？这是从 Python wrapper 迁移到 strategy-orchestrator 的核心架构问题。")

heading(doc, "六、可请对方给出的具体产出")
bullet(doc, "一套市场战略 Agent 的评估指标体系。")
bullet(doc, "一个复杂任务自主编排的参考架构图或流程。")
bullet(doc, "RAG 知识库建设和 metadata 规范建议。")
bullet(doc, "结构化数据 + 文档证据融合的冲突处理方案。")
bullet(doc, "后续可以合作验证的 Demo 场景，例如“某价格带新能源 SUV 市场机会分析”。")

heading(doc, "七、会后建议动作")
bullet(doc, "把对方建议转成 P0/P1/P2 任务清单。")
bullet(doc, "优先修复当前工程问题：Decimal 序列化、StageStatus.warning、文档入库接口、RAG 文档覆盖。")
bullet(doc, "选一个示范场景做闭环：用户问题 -> 数据查询 -> RAG 证据 -> 分析框架 -> 证据账本 -> 报告输出。")
bullet(doc, "建立一套小规模人工评测集，用来验证 Agent 输出质量。")

heading(doc, "八、当前项目一句话定位")
para(
    doc,
    "当前项目不是简单聊天机器人，而是一个面向汽车市场战略分析的垂直决策智能体。它已经具备结构化数据和基础工具底座，"
    "但下一阶段重点应从“能调用工具”升级为“能自主编排、能解释证据、能评估置信度、能支持人类决策”。",
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("生成：汽车市场战略分析师 Agent｜2026-06-17")
set_font(r, size=9, color=(120, 120, 120))

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
