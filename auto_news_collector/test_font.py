"""
测试字体设置
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 测试 SimSun 字体
p1 = doc.add_paragraph()
run1 = p1.add_run("测试宋体标题")
run1.font.name = "SimSun"
run1.font.size = Pt(12)
run1.font.bold = True
# 设置中文字体
rFonts = run1.font.element.get_or_add_rPr().get_or_add_rFonts()
rFonts.set("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}eastAsia", "SimSun")

p2 = doc.add_paragraph()
run2 = p2.add_run("这是正文内容")
run2.font.name = "SimSun"
run2.font.size = Pt(12)
rFonts2 = run2.font.element.get_or_add_rPr().get_or_add_rFonts()
rFonts2.set("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}eastAsia", "SimSun")

doc.save(r"E:\openclaw\workspace\auto_news_collector\output\font_test.docx")
print("Font test saved")
