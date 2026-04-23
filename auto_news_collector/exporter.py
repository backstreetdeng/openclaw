"""
Word文档导出器 - 使用python-docx生成Word报告
"""
import os
from datetime import datetime
from typing import Dict, List

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    Document = None


class WordExporter:
    """Word文档导出器"""

    def __init__(self):
        self.doc = None

    def _set_font(self, run, font_name="SimSun", font_size=Pt(11), bold=False):
        """设置字体（兼容中文）"""
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        # 设置中文字体
        r = run.font.element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}eastAsia",
            font_name
        )

    def export(
        self,
        results: Dict[str, List[Dict]],
        start_date: datetime,
        end_date: datetime,
        output_path: str
    ):
        """
        导出Word文档

        Args:
            results: 采集结果 {领域: [新闻列表]}
            start_date: 开始日期
            end_date: 结束日期
            output_path: 输出路径
        """
        if Document is None:
            raise ImportError("python-docx未安装，请运行: pip install python-docx")

        self.doc = Document()

        # 设置文档样式
        self._setup_styles()

        # 添加标题
        self._add_title(start_date, end_date)

        # 统计信息（排除分类数据和空值）
        news_results = {}
        for k, v in results.items():
            if k.endswith("_分类") or k.endswith("_分领域"):
                continue
            if isinstance(v, list):
                news_results[k] = v
        total = sum(len(v) for v in news_results.values())
        self._add_summary(results, total)

        # 添加各领域内容
        for domain, news_list in results.items():
            if domain.endswith("_分类") or domain.endswith("_分领域"):
                # 跳过内部使用的分类/分领域数据
                continue
            if not isinstance(news_list, list):
                # 跳过非列表项
                continue
            if domain == "企业要闻" and "企业要闻_分类" in results:
                # 企业要闻直接按分类输出，不重复显示
                self._add_domain_section(domain, [])
                chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
                categorized = results["企业要闻_分类"]
                for idx, (category, cat_news) in enumerate(categorized.items()):
                    if isinstance(cat_news, list) and cat_news:
                        self._add_category_section(f"（{chinese_nums[idx]}）{category}", cat_news)
            elif domain == "新技术/新趋势" and "新技术/新趋势_分领域" in results:
                # 新技术/新趋势直接按分领域输出，不重复显示
                self._add_domain_section(domain, [])
                chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
                sub_results = results["新技术/新趋势_分领域"]
                for idx, (category, cat_news) in enumerate(sub_results.items()):
                    if isinstance(cat_news, list) and cat_news:
                        self._add_category_section(f"（{chinese_nums[idx]}）{category}", cat_news)
            elif news_list:
                self._add_domain_section(domain, news_list)

        # 保存文档
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)

    def _setup_styles(self):
        """设置文档样式"""
        from docx.oxml.ns import qn
        # 设置默认字体
        style = self.doc.styles["Normal"]
        style.font.name = "SimSun"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(
            qn('w:eastAsia'), "SimSun"
        )

    def _set_font(self, run, font_name="SimSun", font_size=Pt(12), bold=False):
        """设置字体（兼容中文）"""
        from docx.oxml.ns import qn
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        # 设置中文字体
        r = run.font.element
        rFonts = r.get_or_add_rPr().get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), font_name)

    def _add_title(self, start_date: datetime, end_date: datetime):
        """添加文档标题"""
        # 使用普通段落而不是heading，确保字体设置生效
        title = self.doc.add_paragraph()
        title_run = title.add_run("中国汽车产业资讯简报")
        self._set_font(title_run, "SimSun", Pt(22), bold=True)
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 副标题
        subtitle = self.doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"{start_date.strftime('%Y年%m月%d日')} - {end_date.strftime('%Y年%m月%d日')}")
        self._set_font(subtitle_run, "SimSun", Pt(12))
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()  # 空行

    def _add_summary(self, results: Dict, total: int):
        """添加统计摘要"""
        # 使用普通段落确保字体设置生效
        summary = self.doc.add_paragraph()
        summary_run = summary.add_run("数据概览")
        summary_run.font.name = "SimSun"
        summary_run.font.size = Pt(14)
        summary_run.font.bold = True

        # 计算领域数量
        domain_count = 0
        domain_categories = {}  # 存储分类统计
        
        for k, v in results.items():
            if not v or k.endswith("_分类") or not isinstance(v, list):
                continue
            # 企业要闻的分领域不算独立领域，按分类展示
            if k == "企业要闻" and "企业要闻_分类" in results:
                cat_count = sum(1 for cat_news in results["企业要闻_分类"].values() if cat_news)
                domain_count += cat_count
                continue
            # 新技术/新趋势的分领域不算独立领域，按分领域展示
            if k == "新技术/新趋势" and "新技术/新趋势_分领域" in results:
                cat_count = sum(1 for cat_news in results["新技术/新趋势_分领域"].values() if cat_news)
                domain_count += cat_count
                continue
            domain_count += 1

        p = self.doc.add_paragraph()
        p.add_run(f"共采集 ").font.size = Pt(11)
        bold_run = p.add_run(f"{total}")
        bold_run.bold = True
        bold_run.font.size = Pt(14)
        bold_run.font.color.rgb = RGBColor(0, 102, 204)
        p.add_run(f" 条资讯，涵盖 ").font.size = Pt(11)
        bold_run2 = p.add_run(f"{domain_count}")
        bold_run2.bold = True
        bold_run2.font.size = Pt(14)
        bold_run2.font.color.rgb = RGBColor(0, 102, 204)
        p.add_run(f" 个领域").font.size = Pt(11)

        # 各领域统计
        domain_stats = self.doc.add_paragraph()
        stats_parts = []
        for k, v in results.items():
            if not v or k.endswith("_分类") or not isinstance(v, list):
                continue
            # 企业要闻按分类展示
            if k == "企业要闻" and "企业要闻_分类" in results:
                chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
                for idx, (cat, cat_news) in enumerate(results["企业要闻_分类"].items()):
                    if cat_news:
                        stats_parts.append(f"（{chinese_nums[idx]}）{cat}: {len(cat_news)}条")
            # 新技术/新趋势按分领域展示
            elif k == "新技术/新趋势" and "新技术/新趋势_分领域" in results:
                chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
                for idx, (cat, cat_news) in enumerate(results["新技术/新趋势_分领域"].items()):
                    if cat_news:
                        stats_parts.append(f"（{chinese_nums[idx]}）{cat}: {len(cat_news)}条")
            else:
                stats_parts.append(f"{k}: {len(v)}条")
        stats_text = "  |  ".join(stats_parts)
        domain_stats.add_run(stats_text).font.size = Pt(10)

        self.doc.add_paragraph()  # 空行

    def _add_domain_section(self, domain: str, news_list: List[Dict]):
        """添加领域章节"""
        # 防御性检查
        if not isinstance(news_list, list):
            return
        
        # 领域标题 - 使用普通段落确保字体设置生效
        heading = self.doc.add_paragraph()
        heading_run = heading.add_run(f"【{domain}】")
        self._set_font(heading_run, "SimSun", Pt(14), bold=True)

        # 分割线
        p = self.doc.add_paragraph()
        run = p.add_run("─" * 50)
        run.font.color.rgb = RGBColor(200, 200, 200)
        self._set_font(run, "SimSun")

        for i, news in enumerate(news_list, 1):
            # 防御性检查：确保news是dict
            if not isinstance(news, dict):
                continue
            # 新闻标题
            title_p = self.doc.add_paragraph()
            title_run = title_p.add_run(f"{i}. {news.get('title', '无标题')}")
            self._set_font(title_run, "SimSun", Pt(12), bold=True)

            # 元信息
            meta_p = self.doc.add_paragraph()
            meta_run = meta_p.add_run(f"时间: {news.get('date', '未知日期')}  |  来源: {news.get('source', '未知')}")
            self._set_font(meta_run, "SimSun", Pt(10))
            meta_run.font.color.rgb = RGBColor(128, 128, 128)

            # 链接
            if news.get("link"):
                link_p = self.doc.add_paragraph()
                link_run = link_p.add_run(f"{news.get('link', '')}")
                self._set_font(link_run, "SimSun", Pt(10))
                link_run.font.color.rgb = RGBColor(0, 102, 204)

            # 正文内容
            if news.get("content"):
                content_p = self.doc.add_paragraph()
                content_run = content_p.add_run(news.get("content", ""))
                self._set_font(content_run, "SimSun", Pt(12))

            # 空行
            self.doc.add_paragraph()

        # 章节结束分割线
        end_p = self.doc.add_paragraph()
        end_run = end_p.add_run("─" * 50)
        end_run.font.color.rgb = RGBColor(200, 200, 200)
        self._set_font(end_run, "SimSun")
        self.doc.add_paragraph()  # 空行

    def _add_category_section(self, category: str, news_list: List[Dict]):
        """添加分领域章节（用于企业要闻的子分类）"""
        # 防御性检查
        if not isinstance(news_list, list):
            return
        # 子分类标题 - 使用普通段落确保字体设置生效
        heading = self.doc.add_paragraph()
        heading_run = heading.add_run(f"  {category}")
        self._set_font(heading_run, "SimSun", Pt(12), bold=True)

        for i, news in enumerate(news_list, 1):
            # 新闻标题
            title_p = self.doc.add_paragraph()
            title_run = title_p.add_run(f"{i}. {news.get('title', '无标题')}")
            self._set_font(title_run, "SimSun", Pt(12), bold=True)

            # 元信息
            meta_p = self.doc.add_paragraph()
            meta_run = meta_p.add_run(f"时间: {news.get('date', '未知日期')}  |  来源: {news.get('source', '未知')}")
            self._set_font(meta_run, "SimSun", Pt(10))
            meta_run.font.color.rgb = RGBColor(128, 128, 128)

            # 链接
            if news.get("link"):
                link_p = self.doc.add_paragraph()
                link_run = link_p.add_run(f"{news.get('link', '')}")
                self._set_font(link_run, "SimSun", Pt(10))
                link_run.font.color.rgb = RGBColor(0, 102, 204)

            # 正文内容
            if news.get("content"):
                content_p = self.doc.add_paragraph()
                content_run = content_p.add_run(news.get("content", ""))
                self._set_font(content_run, "SimSun", Pt(12))

            # 空行
            self.doc.add_paragraph()
